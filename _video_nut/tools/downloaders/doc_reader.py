#!/usr/bin/env python3
"""
Word Document (.docx) Reader and Keyword Searcher

Downloads and parses Word document files (.docx) from a local path or URL.
Uses python-docx as the primary engine, and falls back to a custom atomic ZIP/XML 
parser if python-docx is not installed on the system.

Usage:
    python doc_reader.py --url "https://example.com/report.docx"
    python doc_reader.py --path "./report.docx" --search "scam"
    python doc_reader.py --url "https://example.com/report.docx" --output "./Projects/test/assets/doc.docx"
"""

import sys
import os
import argparse
import requests
import io
import time
from random import uniform
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

# Enforce UTF-8 output encoding for Windows terminal safety
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')

# Set path for importing audit_logger
vn_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(vn_root))

try:
    from tools.logging import audit_logger
    has_audit_logger = True
except ImportError:
    has_audit_logger = False

# Try to import python-docx
try:
    import docx
    has_docx = True
except ImportError:
    has_docx = False


def log_action_to_audit(project_path, action, url="", local_path="", status="ok", details=""):
    """Wrapper for logging to audit trail if available."""
    if has_audit_logger:
        audit_logger.log_action(
            project_path=project_path,
            category="read",
            action=action,
            url=url,
            local_path=local_path,
            status=status,
            details=details
        )


def read_docx_raw_xml(file_path_or_bytes) -> str:
    """
    Fallback parser that extracts text from .docx file without external dependencies
    by unzipping and parsing the main word/document.xml file.
    """
    paragraphs = []
    try:
        # docx file is actually a zip file
        with zipfile.ZipFile(file_path_or_bytes) as docx_zip:
            xml_content = docx_zip.read('word/document.xml')
            root = ET.fromstring(xml_content)
            
            # Extract paragraphs by searching for <w:p> elements
            # Using endswith to bypass namespaces
            for child in root.iter():
                if child.tag.endswith('}p') or child.tag == 'p':
                    p_text = []
                    for t_elem in child.iter():
                        if t_elem.tag.endswith('}t') or t_elem.tag == 't':
                            if t_elem.text:
                                p_text.append(t_elem.text)
                    paragraphs.append("".join(p_text))
                    
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[Error in raw XML fallback parser] {e}"


def extract_docx_text(file_path_or_bytes) -> tuple[str, list[dict], str]:
    """
    Parses paragraphs and tables using python-docx if available, 
    otherwise falls back to the XML zip parser.
    
    Returns:
        tuple containing (full_text, tables_list_of_dicts, engine_used)
    """
    if has_docx:
        try:
            doc = docx.Document(file_path_or_bytes)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
                
            # Extract table text
            tables_data = []
            for i, table in enumerate(doc.tables):
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                tables_data.append({
                    "table_index": i + 1,
                    "content": "\n".join(table_text)
                })
                full_text.append(f"\n--- [Table {i+1}] ---\n" + "\n".join(table_text))
                
            return "\n".join(full_text), tables_data, "python-docx"
        except Exception as e:
            # If python-docx fails for some reason, try the raw XML parser
            xml_text = read_docx_raw_xml(file_path_or_bytes)
            return xml_text, [], f"raw-xml-fallback (python-docx failed: {e})"
    else:
        xml_text = read_docx_raw_xml(file_path_or_bytes)
        return xml_text, [], "raw-xml-fallback"


def download_and_read_docx(url=None, file_path=None, search_term=None, output_path=None, project_dir=None):
    """
    Retrieves and reads a Word Document.
    """
    if not url and not file_path:
        print("Error: Either --url or --path must be specified.")
        sys.exit(1)

    # Determine project folder for logging
    if not project_dir:
        if output_path:
            project_dir = str(Path(output_path).parent.parent)
        else:
            project_dir = "."

    doc_data = None
    source_desc = ""

    # 1. Download file from URL if provided
    if url:
        source_desc = url
        # Add rate limiting wait delay
        delay = uniform(1, 2)
        print(f"Waiting {delay:.2f} seconds before downloading {url}...")
        time.sleep(delay)

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }

        try:
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            doc_data = io.BytesIO(response.content)
            
            # Save downloaded file
            if output_path:
                out_path = Path(output_path)
                out_path.parent.mkdir(parents=True, exist_ok=True)
                
                # Write to temp file then rename (atomic)
                temp_path = out_path.with_suffix(".tmp")
                with open(temp_path, "wb") as f:
                    f.write(response.content)
                if temp_path.stat().st_size == 0:
                    raise IOError("Downloaded file is 0 bytes")
                temp_path.replace(out_path)
                
                print(f"[SAVE] Word document downloaded and saved to: {out_path}")
                log_action_to_audit(project_dir, "Downloaded DOCX file", url=url, local_path=str(out_path), status="ok")
        except Exception as e:
            print(f"[FAIL] Error downloading DOCX: {e}")
            log_action_to_audit(project_dir, f"Failed to download DOCX: {e}", url=url, status="failed", details=str(e))
            sys.exit(1)
            
    # 2. Read from local path
    elif file_path:
        source_desc = file_path
        local_file = Path(file_path)
        if not local_file.exists():
            print(f"Error: Local file {file_path} does not exist.")
            sys.exit(1)
        doc_data = file_path
        log_action_to_audit(project_dir, f"Reading local DOCX file", local_path=str(local_file), status="ok")

    # 3. Extract Text
    try:
        full_text, tables, engine = extract_docx_text(doc_data)
        lines = full_text.split("\n")
        print(f"[DOC] DOCX loaded using {engine}: {len(lines)} lines found")
        
        # 4. Search keyword if specified
        if search_term:
            print(f"[SCAN] Searching for: '{search_term}'")
            search_lower = search_term.lower()
            matches = []
            
            for idx, line in enumerate(lines):
                if search_lower in line.lower():
                    context_start = max(0, idx - 3)
                    context_end = min(len(lines), idx + 4)
                    context = "\n".join(lines[context_start:context_end])
                    matches.append({
                        "line_num": idx + 1,
                        "line": line.strip(),
                        "context": context
                    })
                    
            if matches:
                print(f"\n[OK] Found {len(matches)} matches for '{search_term}':\n")
                for m_idx, match in enumerate(matches[:15]):
                    print(f"{'='*60}")
                    print(f"[MATCH] Match {m_idx+1} - Line {match['line_num']}")
                    print(f"{'='*60}")
                    print(f"Line: {match['line']}")
                    print(f"\nContext:")
                    print(match['context'])
                    print()
                if len(matches) > 15:
                    print(f"... and {len(matches) - 15} more matches")
                
                log_action_to_audit(
                    project_dir,
                    f"Searched DOCX for '{search_term}' (Found {len(matches)} matches)",
                    url=url or "",
                    local_path=output_path or file_path or "",
                    status="ok",
                    details=f"Engine: {engine}"
                )
            else:
                print(f"[FAIL] No matches found for '{search_term}'")
                log_action_to_audit(
                    project_dir,
                    f"Searched DOCX for '{search_term}' (0 matches)",
                    url=url or "",
                    local_path=output_path or file_path or "",
                    status="skipped"
                )
            return

        # 5. Output default content
        # Limit to reasonable screen print length
        max_chars = 15000
        if len(full_text) > max_chars:
            intro = full_text[:7000]
            outro = full_text[-7000:]
            truncated = len(full_text) - max_chars
            print(f"--- DOCX CONTENT START ---")
            print(intro)
            print(f"\n\n[... {truncated:,} characters truncated from middle ...]\n\n")
            print(f"--- DOCX CONTENT END ---")
            print(outro)
        else:
            print(full_text)
            
        log_action_to_audit(
            project_dir,
            f"Read DOCX file content ({len(full_text)} chars)",
            url=url or "",
            local_path=output_path or file_path or "",
            status="ok",
            details=f"Engine: {engine}"
        )
            
    except Exception as e:
        print(f"[FAIL] Error parsing DOCX: {e}")
        log_action_to_audit(
            project_dir,
            f"Failed to parse DOCX: {e}",
            url=url or "",
            local_path=output_path or file_path or "",
            status="failed",
            details=str(e)
        )
        sys.exit(1)


def main():
    parser = argparse.ArgumentParser(
        description="Download, read, and search Microsoft Word (.docx) files.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Download and print DOCX contents
  python doc_reader.py --url "https://example.com/brief.docx"
  
  # Read and search local file
  python doc_reader.py --path "./Projects/report.docx" --search "scandal"
  
  # Download, save locally, and search
  python doc_reader.py --url "https://example.com/brief.docx" --output "./Projects/test/doc.docx" --search "budget"
        """
    )
    
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--url", help="URL of the Word document to download")
    group.add_argument("--path", help="Local path of the Word document to read")
    
    parser.add_argument("--search", "-s", help="Search keyword inside the document")
    parser.add_argument("--output", "-o", help="Path to save the downloaded document")
    parser.add_argument("--project-dir", help="Project directory path for logging")
    
    args = parser.parse_args()
    
    download_and_read_docx(
        url=args.url,
        file_path=args.path,
        search_term=args.search,
        output_path=args.output,
        project_dir=args.project_dir
    )


if __name__ == "__main__":
    main()
